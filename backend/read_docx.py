from docx import Document
import sys

def read_docx(file_path):
    doc = Document(file_path)
    for i, para in enumerate(doc.paragraphs):
        print(f"[{i}]: {para.text}")
    
    print("\n--- TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            print(f"  Row {r_idx}: {row_data}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
